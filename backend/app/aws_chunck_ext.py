import re
from datetime import datetime
from langchain_aws import BedrockEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
import docx
import boto3
import os
import io

embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v2:0")

def read_docx(file_path):
    """Extract text from a .docx file."""
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(file_path):
    """
    Extract text from a .pdf file using AWS Textract.
    Supports local PDF files.
    """

    textract_client = boto3.client(
        "textract",
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
        region_name=os.getenv("AWS_REGION")
    )

    with open(file_path, "rb") as f:
        pdf_bytes = f.read()

    response = textract_client.analyze_document(
        Document={"Bytes": pdf_bytes},
        FeatureTypes=["TABLES", "FORMS"]
    )

    text = ""
    for block in response["Blocks"]:
        if block["BlockType"] == "LINE":
            text += block["Text"] + "\n"

    return text

def read_document(file_path):
    """Read a document and extract text content from .docx or .pdf."""
    if file_path.endswith(".docx"):
        return read_docx(file_path)
    elif file_path.endswith(".pdf"):
        return read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type. Please provide a .docx or .pdf file.")


def split_into_chunks(content, chunk_size=3000, overlap=500):
    """Split text into overlapping chunks."""
    return [content[i:i + chunk_size] for i in range(0, len(content), chunk_size - overlap)]

def perform_semantic_search(chunks, query, k=15):
    """Retrieve top-k semantically relevant chunks for the given query."""
    documents = [Document(page_content=chunk) for chunk in chunks]
    vector_store = FAISS.from_documents(documents, embeddings)
    results = vector_store.similarity_search(query, k=k)
    return [result.page_content for result in results]

date_patterns = [
    r'\b(20\d{2}|19\d{2})\b',
    r'\b(0[1-9]|1[0-2])/20\d{2}\b',
    r'\b(0[1-9]|[12][0-9]|3[01])/(0[1-9]|1[0-2])/20\d{2}\b'
]
combined_date_pattern = '|'.join(date_patterns)

work_certification_keywords = [
    "experience", "worked", "employment", "job", "position", "role", "responsibility",
    "certified", "certificate", "certification", "completed", "course", "training", "accreditation"
]

def contains_relevant_keywords(chunk, keywords):
    """Check if the chunk contains relevant keywords."""
    return any(keyword.lower() in chunk.lower() for keyword in keywords)

def filter_recent_chunks(chunks, years=5):
    """Keep only chunks that mention recent (last n years) experience or certification."""
    current_year = datetime.now().year
    filtered_chunks = []

    for chunk in chunks:
        date_matches = re.findall(combined_date_pattern, chunk)
        if contains_relevant_keywords(chunk, work_certification_keywords):
            for match in date_matches:
                if isinstance(match, tuple):
                    match = next((m for m in match if m), None)
                if not match:
                    continue

                year = None
                if re.match(r'^\d{4}$', match):
                    year = int(match)
                elif re.search(r'/(\d{4})', match):
                    year = int(re.search(r'/(\d{4})', match).group(1))

                if year and (current_year - year) <= years:
                    filtered_chunks.append(chunk)
                    break

    return filtered_chunks

def find_relevant_experience(file_path, query, chunk_size=2000, overlap=500, top_k=10, years=5):
    """Perform full pipeline: read, chunk, search, and filter resume content."""
    document_content = read_document(file_path)
    chunks = split_into_chunks(document_content, chunk_size, overlap)
    top_chunks = perform_semantic_search(chunks, query, k=top_k)
    recent_chunks = filter_recent_chunks(top_chunks, years)
    first_chunk = chunks[0] if chunks else ""
    return recent_chunks, first_chunk

def final_chunks(file_path):
    """Wrapper for Streamlit app to return relevant experience chunks."""
    query = "Retrieve relevant work experience or certifications from the last 5 years"
    recent_chunks, first_chunk=find_relevant_experience(file_path, query)
    return {"experience_chunks": recent_chunks, "header_chunk": first_chunk}

def header_chunks(file_path):
    """Wrapper for Streamlit app to return all chunks for header processing."""
   
