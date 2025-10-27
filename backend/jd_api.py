import os
import re
import boto3
import requests
import base64
from io import BytesIO
from flask import Flask, request, jsonify
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from langchain_aws import ChatBedrock

# ================== FLASK & ENV CONFIG ==================
app = Flask(__name__)
load_dotenv()

AWS_ACCESS_KEY = os.getenv("access_key")
AWS_SECRET_KEY = os.getenv("secret_access_key")
AWS_REGION = os.getenv("region_name", "ap-south-1")
SERP_API_KEY = os.getenv("SERP_API_KEY")

session = boto3.Session(
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)

llm = ChatBedrock(
    model_id="anthropic.claude-3-sonnet-20240229-v1:0",
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region=AWS_REGION
)

# ================== HELPERS ==================

def clean_bedrock_response(response) -> str:
    text = response if isinstance(response, str) else str(response)
    match = re.search(r'content="(.*?)"\s*(?:additional_kwargs|response_metadata|$)', text, re.DOTALL)
    clean = match.group(1) if match else text
    clean = clean.encode('utf-8').decode('unicode_escape')
    clean = re.sub(r'additional_kwargs=.*', '', clean)
    clean = re.sub(r'response_metadata=.*', '', clean)
    clean = re.sub(r'id=.*', '', clean)
    clean = re.sub(r'usage_metadata=.*', '', clean)
    return clean.strip().strip('"').strip("'").strip()


def ask_bedrock(prompt: str) -> str:
    try:
        response = llm.invoke(input=prompt)
        return clean_bedrock_response(response)
    except Exception as e:
        return f"Bedrock error: {e}"


def create_word_doc(content: str, filename: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Job Description", level=1)

    sections = re.split(
        r'^(Job Description: AI Developer|Intro Section:|Objectives of this role:|Your tasks:|Required skills and qualifications:|Preferred skills and qualifications:)',
        content, flags=re.MULTILINE | re.IGNORECASE
    )

    for i in range(1, len(sections), 2):
        title = sections[i].strip()
        body = sections[i + 1].strip()

        if not body:
            continue

        if title.lower().startswith(('intro section', 'objectives', 'your tasks', 'required', 'preferred')):
            doc.add_heading(title, level=2)
        elif title.lower().startswith('job description'):
            intro_lines = body.split("\n\n")[0].split("\n")
            for line in intro_lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(8)
            continue

        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.lstrip().startswith(('•', '-', '*')):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(('•', '-', '*')):
                        cleaned_line = re.sub(r'^[•\-\*\s\t]+', '', line).strip()
                        if cleaned_line:
                            doc.add_paragraph(cleaned_line, style="List Bullet")
            else:
                p = doc.add_paragraph(block)
                p.paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ================== FLASK ENDPOINT ==================

@app.route("/generate_jd", methods=["POST"])
def generate_jd():
    """
    One-shot endpoint that:
    1. Accepts requirements JSON
    2. Fetches market JDs
    3. Extracts insights
    4. Generates JD with Claude
    5. Checks bias
    6. Returns JD text, feedback, and Base64 Word file
    """
    try:
        data = request.get_json()
        role = data.get("role")
        experience = data.get("experience", "")
        domain = data.get("domain", "")
        location = data.get("location", "")
        job_responsibilities = data.get("job_responsibilities", "")
        company_values = data.get("company_values", "")
        diversity = data.get("diversity", "")

        if not role:
            return jsonify({"error": "Missing required field: 'role'"}), 400

        # --- Step 1: Fetch Market JDs ---
        url = "https://serpapi.com/search"
        params = {
            "engine": "google_jobs",
            "q": role,
            "hl": "en",
            "num": 5,
            "api_key": SERP_API_KEY
        }

        try:
            resp = requests.get(url, params=params)
            resp.raise_for_status()
            jobs = resp.json().get("jobs_results", [])
            postings = [job.get("description", "") for job in jobs]
        except Exception as e:
            postings = []
            print(f"SERP error: {e}")

        # --- Step 2: Analyze Market Data ---
        if postings:
            text = "\n\n".join(postings[:5])
            analysis_prompt = f"""
            You are an HR analyst. Analyze the following job descriptions:

            {text}

            1. List trending technical and soft skills.
            2. Highlight frequently used phrases.
            3. Identify bias-prone or exclusive terms (e.g., 'rockstar', 'young').
            """
            analysis = ask_bedrock(analysis_prompt)
        else:
            analysis = "No external market data available."

        # --- Step 3: Generate JD ---
        jd_prompt = f"""
        Draft an inclusive, engaging, SEO-optimized Job Description in the following template format.

        ============================
        Job Description: Template
        We are looking for a passionate {role} to join our team at [Company X], located in {location}.

        Intro Section:
        Provide 1–2 paragraphs introducing the role, responsibilities, and why this job is exciting.
        Explicitly mention the key project(s) or initiatives the new hire will be working on, based on:
        "{job_responsibilities}"

        Objectives of this role:
        - List 4–5 key objectives directly connected to these projects and the role’s impact.

        Your tasks:
        - List 4–6 detailed tasks reflecting actual day-to-day work related to the provided project(s).
        - Ensure they align with the overall goals mentioned above.

        Required skills and qualifications:
        - 5–7 bullet points with mandatory skills, education, or certifications.
        - Include both technical and soft skills.
        - Include years of experience if applicable ({experience}).

        Preferred skills and qualifications:
        - 3–5 bullet points with additional desirable or domain-specific skills.

        ============================

        Requirements:
        Role: {role}
        Experience: {experience}
        Domain: {domain}
        Location: {location}
        Company Values: {company_values}
        Diversity Considerations: {diversity}
        Key Projects: {job_responsibilities}

        Market Insights:
        {analysis}

        Instructions:
        - Follow the template exactly.
        - Use inclusive, bias-free, gender-neutral language.
        - Naturally embed relevant keywords from market insights.
        - Include the location naturally in the introduction and conclusion.
        - End with a call-to-action encouraging candidates to apply.
        """
        jd_text = ask_bedrock(jd_prompt)

        # --- Step 4: Bias Review ---
        bias_prompt = f"""
        Review the following Job Description for biased, non-inclusive, or exclusionary language:

        {jd_text}

        Suggest improvements in bullet points.
        """
        bias_feedback = ask_bedrock(bias_prompt)

        # --- Step 5: Generate Word File ---
        buffer = create_word_doc(jd_text, f"{role.replace(' ', '_')}_JD.docx")
        base64_doc = base64.b64encode(buffer.getvalue()).decode("utf-8")

        # --- Step 6: Return Response ---
        return jsonify({
            "role": role,
            "analysis": analysis,
            "job_description": jd_text,
            "bias_feedback": bias_feedback,
            "word_doc_base64": base64_doc,
            "filename": f"{role.replace(' ', '_')}_JD.docx"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ================== MAIN ==================
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
