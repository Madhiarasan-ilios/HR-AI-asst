# ============================
#      MERGED FLASK APP
# ============================

import os
import re
import io
import csv
import base64
import tempfile
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from typing import List, Dict, Any

from flask import Flask, request, jsonify
from flask_cors import CORS
import pandas as pd
import requests
import boto3
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# ---- Bedrock + AWS imports from Code2 ----
import asyncio
import logging
from aws_services.polly import PollyClient
from aws_services.transcribe import start_transcription
from langchain_logic.chains import (
    get_question_generation_chain,
    get_report_generation_chain,
    get_short_answer_evaluation_chain,
    get_structured_extraction_chain,
)

# ---- LangChain AWS models (Code1) ----
from langchain_aws import ChatBedrock, BedrockLLM

# ---- Resume evaluation from Code1 ----
from app.aws_parsing import evaluate_resume_skills_with_time, calculate_relevance
from app.aws_skillset import final_claude
from app.aws_chunck_ext import final_chunks

# ============================
#       INITIAL SETUP
# ============================

load_dotenv()

app = Flask(__name__)
CORS(app)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

AWS_ACCESS_KEY = os.getenv("access_key")
AWS_SECRET_KEY = os.getenv("secret_access_key")
AWS_REGION = os.getenv("region_name", "ap-south-1")
SERP_API_KEY = os.getenv("SERP_API_KEY")

REGION = AWS_REGION
MODEL_ID = "anthropic.claude-3-sonnet-20240229-v1:0"   # interviewer model

# ============================
#     AWS SESSIONS & LLMs
# ============================

session = boto3.Session(
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)

claude_llm = ChatBedrock(
    model_id="anthropic.claude-3-sonnet-20240229-v1:0",
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region=AWS_REGION
)

llama_llm = BedrockLLM(
    model_id="meta.llama3-70b-instruct-v1:0",
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region=AWS_REGION
)

# ============================
#      INTERVIEW SESSION MEMORY
# ============================

SESSION_STORE: Dict[str, Dict[str, Any]] = {}


def init_session(session_id: str = "active"):
    """
    Initialize or reset an interview session.
    For now we use a single 'active' session ID.
    """
    SESSION_STORE[session_id] = {
        "current_question_index": 0,
        "questions": [],
        "transcript": [],  # list of {question, answer}
        "structured": {
            "name": "",
            "experience": "",
            "current_ctc": "",
            "expected_ctc": "",
            "notice_period": "",
            "skills": [],
            "project_highlight": ""
        }
    }


# ============================
#      SHARED HELPERS
# ============================

def clean_bedrock_response(response) -> str:
    text = response if isinstance(response, str) else str(response)
    match = re.search(
        r'content="(.*?)"\s*(?:additional_kwargs|response_metadata|$)',
        text,
        re.DOTALL
    )
    clean = match.group(1) if match else text
    clean = clean.encode('utf-8').decode('unicode_escape')
    clean = re.sub(r'additional_kwargs=.*', '', clean)
    clean = re.sub(r'response_metadata=.*', '', clean)
    clean = re.sub(r'id=.*', '', clean)
    clean = re.sub(r'usage_metadata=.*', '', clean)
    return clean.strip().strip('"').strip("'").strip()


def ask_bedrock(prompt: str, use_model="claude") -> str:
    try:
        llm = claude_llm if use_model == "claude" else llama_llm
        response = llm.invoke(input=prompt)
        return clean_bedrock_response(response)
    except Exception as e:
        return f"Bedrock error: {e}"


def create_word_doc(content: str, filename: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Job Description", level=1)
    sections = re.split(
        r'^(Job Description:|Intro Section:|Objectives of this role:|Your tasks:|Required skills and qualifications:|Preferred skills and qualifications:)',
        content,
        flags=re.MULTILINE | re.IGNORECASE
    )

    for i in range(1, len(sections), 2):
        title = sections[i].strip()
        body = sections[i + 1].strip()
        if not body:
            continue
        if title.lower().startswith(('intro', 'objectives', 'your tasks', 'required', 'preferred')):
            doc.add_heading(title, level=2)
        elif title.lower().startswith('job description'):
            for line in body.splitlines():
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.space_after = Pt(8)
            continue
        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith(('•', '-', '*')):
                for line in block.splitlines():
                    cleaned_line = re.sub(r'^[•\-\*\s\t]+', '', line).strip()
                    if cleaned_line:
                        doc.add_paragraph(cleaned_line, style="List Bullet")
            else:
                p = doc.add_paragraph(block)
                p.paragraph_format.space_after = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def save_upload_to_temp(upload_file) -> str:
    suffix = os.path.splitext(upload_file.filename or "")[1]
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tf.write(upload_file.read())
    tf.flush()
    tf.close()
    return tf.name


# ============================
#     ENDPOINT 1: GENERATE JD
# ============================

@app.route("/generate_jd", methods=["POST"])
def generate_jd():
    try:
        data = request.get_json()
        role = data.get("role")
        experience = data.get("experience", "")
        domain = data.get("domain", "")
        location = data.get("location", "")
        job_responsibilities = data.get("job_responsibilities", "")
        company_values = data.get("company_values", "")
        diversity = data.get("diversity", "")

        if not role:
            return jsonify({"error": "Missing required field: 'role'"}), 400

        postings = []
        try:
            resp = requests.get("https://serpapi.com/search", params={
                "engine": "google_jobs",
                "q": role,
                "hl": "en",
                "num": 5,
                "api_key": SERP_API_KEY
            })
            resp.raise_for_status()
            jobs = resp.json().get("jobs_results", [])
            postings = [j.get("description", "") for j in jobs]
        except Exception as e:
            print(f"SERP error: {e}")

        if postings:
            joined = "\n\n".join(postings[:5])
            analysis_prompt = f"""
            Analyze the following job descriptions and summarize:
            1. Trending technical and soft skills
            2. Common phrases
            3. Biased or exclusive terms

            {joined}
            """
            analysis = ask_bedrock(analysis_prompt, use_model="claude")
        else:
            analysis = "No market data found."

        jd_prompt = f"""
        Write an inclusive Job Description for {role} in {location}.
        Experience: {experience}
        Domain: {domain}
        Responsibilities: {job_responsibilities}
        Company Values: {company_values}
        Diversity: {diversity}

        Market Insights:
        {analysis}

        Follow this structure:
        Job Description:
        Intro Section:
        Objectives of this role:
        Your tasks:
        Required skills and qualifications:
        Preferred skills and qualifications:
        """
        jd_text = ask_bedrock(jd_prompt, use_model="claude")

        bias_prompt = f"Review this JD for bias and suggest improvements:\n\n{jd_text}"
        bias_feedback = ask_bedrock(bias_prompt, use_model="claude")

        buffer = create_word_doc(jd_text, f"{role}_JD.docx")
        b64_doc = base64.b64encode(buffer.getvalue()).decode()

        return jsonify({
            "role": role,
            "analysis": analysis,
            "job_description": jd_text,
            "bias_feedback": bias_feedback,
            "word_doc_base64": b64_doc,
            "filename": f"{role.replace(' ', '_')}_JD.docx"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============================
#  ENDPOINT 2: SOURCE CANDIDATES
# ============================

FEW_SHOT_EXAMPLES = [
    {
        "input": "Full Stack Developer with React, Node.js and MySQL",
        "output": '("Full Stack Developer" OR "Fullstack Engineer") AND (React OR Angular OR Vue) AND (Node.js OR Django) AND (MySQL OR MongoDB)'
    },
    {
        "input": "Entry level developer, exclude senior",
        "output": '("junior developer" OR "associate developer") ("1 year" OR "entry level") -senior -lead'
    }
]


@app.route("/source_candidates", methods=["POST"])
def source_candidates():
    try:
        data = request.get_json()
        jd_text = data.get("job_description", "")
        num_results = int(data.get("num_results", 10))
        if not jd_text:
            return jsonify({"error": "Missing job_description"}), 400

        examples_text = "\n\n".join([
            f"Job Description: {ex['input']}\nBoolean Query: {ex['output']}"
            for ex in FEW_SHOT_EXAMPLES
        ])

        prompt = f"""
        Convert the job description into a LinkedIn Boolean query.

        {examples_text}

        Job Description: {jd_text}
        Boolean Query:
        """

        boolean_query = ask_bedrock(prompt, use_model="llama")

        url = "https://serpapi.com/search"
        params = {
            "engine": "google",
            "q": f"site:linkedin.com/in/ {boolean_query}",
            "num": num_results,
            "api_key": SERP_API_KEY
        }
        resp = requests.get(url, params=params)
        data = resp.json()

        profiles = [
            {"Name": r.get("title"), "Snippet": r.get("snippet"), "Link": r.get("link")}
            for r in data.get("organic_results", [])
        ]

        df = pd.DataFrame(profiles)
        csv_buf = io.StringIO()
        df.to_csv(csv_buf, index=False)
        csv_b64 = base64.b64encode(csv_buf.getvalue().encode()).decode()

        return jsonify({
            "boolean_query": f"site:linkedin.com/in/ {boolean_query}",
            "profiles_found": len(profiles),
            "profiles": profiles,
            "candidates_csv_base64": csv_b64,
            "filename": "candidates.csv"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============================
#  ENDPOINT 3: RESUME EVALUATION
# ============================

@app.route("/evaluate_resumes", methods=["POST"])
def evaluate_resumes():
    try:
        start_time = datetime.now()
        jd_file = request.files.get("jd_file")
        resumes = request.files.getlist("resumes")

        if not jd_file or not resumes:
            return jsonify({"error": "Missing jd_file or resumes"}), 400

        jd_path = save_upload_to_temp(jd_file)
        ranking_skills = final_claude(jd_path)

        results = []
        temp_files = [jd_path]

        def process_resume(resume_file):
            try:
                path = save_upload_to_temp(resume_file)
                temp_files.append(path)
                doc_content = final_chunks(path)
                weights_json, in_tokens, out_tokens = evaluate_resume_skills_with_time(
                    llama_llm, doc_content, ranking_skills
                )
                score, reasons = calculate_relevance(weights_json)
                cost = (in_tokens * 0.00318) + (out_tokens * 0.0042)
                return {
                    "Candidate Name": os.path.splitext(resume_file.filename)[0],
                    "Similarity (%)": round(score, 2),
                    "Reasons": reasons,
                    "Cost": round(cost, 4)
                }
            except Exception as e:
                return {
                    "Candidate Name": resume_file.filename,
                    "Similarity (%)": 0,
                    "Reasons": str(e),
                    "Cost": 0
                }

        with ThreadPoolExecutor() as pool:
            for f in pool.map(process_resume, resumes):
                results.append(f)

        df = pd.DataFrame(results).sort_values(by="Similarity (%)", ascending=False)
        csv_data = df.to_csv(index=False)
        csv_b64 = base64.b64encode(csv_data.encode()).decode()
        duration = (datetime.now() - start_time).total_seconds()

        return jsonify({
            "results": df.to_dict(orient="records"),
            "csv_base64": csv_b64,
            "filename": "results.csv",
            "top_5": df.head(5).to_dict(orient="records"),
            "processing_time_seconds": duration
        })

    finally:
        for f in locals().get("temp_files", []):
            try:
                os.remove(f)
            except:
                pass


# ============================
#     INTERVIEWER ENDPOINTS
# ============================

@app.route("/generate-questions", methods=["POST"])
def generate_questions():
    jd_file = request.files.get("jd")
    resume_file = request.files.get("resume")

    if not jd_file or not resume_file:
        return jsonify({"error": "jd and resume files are required"}), 400

    # Convert files to text
    jd_text = jd_file.read().decode("utf-8", errors="ignore")
    resume_text = resume_file.read().decode("utf-8", errors="ignore")

    chain = get_question_generation_chain(model_id=MODEL_ID, region_name=REGION)
    result = chain.invoke({"jd": jd_text, "resume": resume_text})

    # NEW: initialize interview session and store questions
    init_session("active")
    SESSION_STORE["active"]["questions"] = result.questions

    return jsonify({"questions": result.questions})


@app.route("/speak-question", methods=["POST"])
def speak_question():
    data = request.get_json()
    question = data.get("question")

    if not question:
        return jsonify({"error": "question is required"}), 400

    polly = PollyClient(region_name=REGION)
    audio_stream = polly.synthesize_to_stream(question)
    audio_bytes = audio_stream.read()

    return audio_bytes, 200, {"Content-Type": "audio/mpeg"}


# ============================
#  UPDATED: TRANSCRIBE AUDIO
# ============================

@app.route("/transcribe-audio", methods=["POST"])
def transcribe_audio():
    if "audio" not in request.files:
        return jsonify({"error": "No audio file provided"}), 400

    audio_file = request.files["audio"]
    file_path = save_upload_to_temp(audio_file)

    try:
        # Upload to S3
        s3 = boto3.client("s3", region_name=REGION)
        bucket = os.getenv("TRANSCRIBE_BUCKET_NAME")
        if not bucket:
            return jsonify({"error": "TRANSCRIBE_BUCKET_NAME not set in .env"}), 500

        object_key = os.path.basename(file_path)
        s3.upload_file(file_path, bucket, object_key)

        # Start Transcribe Batch Job
        transcribe = boto3.client("transcribe", region_name=REGION)
        job_name = f"transcribe_job_{int(datetime.now().timestamp())}"

        media_format = os.path.splitext(object_key)[1][1:].lower()
        if media_format not in ["wav", "mp3", "mp4", "flac", "ogg"]:
            media_format = "mp3"  # fallback

        transcribe.start_transcription_job(
            TranscriptionJobName=job_name,
            Media={"MediaFileUri": f"s3://{bucket}/{object_key}"},
            MediaFormat=media_format,
            LanguageCode="en-US",
        )

        return jsonify({
            "message": "Transcription job started",
            "job_name": job_name,
            "audio_file": object_key
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        try:
            os.remove(file_path)
        except:
            pass


# NEW: POLL TRANSCRIPTION RESULT (ENHANCED FLOW)
@app.route("/get-transcription-result", methods=["GET"])
def get_transcription_result():
    data = request.get_json()
    job_name = data.get("job_name") if data else None
    if not job_name:
        return jsonify({"error": "job_name is required"}), 400

    transcribe = boto3.client("transcribe", region_name=REGION)
    response = transcribe.get_transcription_job(
        TranscriptionJobName=job_name
    )
    status = response["TranscriptionJob"]["TranscriptionJobStatus"]

    if status == "FAILED":
        return jsonify({"status": "FAILED"}), 500

    if status != "COMPLETED":
        return jsonify({"status": status}), 200

    transcript_url = response["TranscriptionJob"]["Transcript"]["TranscriptFileUri"]
    transcript_resp = requests.get(transcript_url).json()
    text = transcript_resp["results"]["transcripts"][0]["transcript"]

    # If no active session, just return baseline behavior
    if "active" not in SESSION_STORE or not SESSION_STORE["active"]["questions"]:
        return jsonify({
            "status": "COMPLETED",
            "transcript": text
        })

    state = SESSION_STORE["active"]

    # Store Q&A in transcript
    current_idx = state["current_question_index"]
    if current_idx < len(state["questions"]):
        current_question = state["questions"][current_idx]
    else:
        current_question = ""

    state["transcript"].append({
        "question": current_question,
        "answer": text
    })

    # Short answer evaluation
    try:
        short_chain = get_short_answer_evaluation_chain(model_id=MODEL_ID, region_name=REGION)
        evaluation = short_chain.invoke({"answer": text})
    except Exception as e:
        logger.error(f"Short answer evaluation error: {e}")
        evaluation = None

    if evaluation and getattr(evaluation, "needs_more_detail", False):
        return jsonify({
            "status": "COMPLETED",
            "transcript": text,
            "action": "elaborate",
            "next_question": "Could you please elaborate more on that?"
        })

    # Structured extraction
    try:
        extraction_chain = get_structured_extraction_chain(model_id=MODEL_ID, region_name=REGION)
        structured = extraction_chain.invoke({"answer": text})

        for key, value in structured.model_dump().items():
            if value:
                if isinstance(value, list):
                    existing = state["structured"].get(key, [])
                    state["structured"][key] = list({*existing, *value})
                else:
                    state["structured"][key] = value
    except Exception as e:
        logger.error(f"Structured extraction error: {e}")

    # Move to next question
    state["current_question_index"] += 1

    # If finished all questions
    if state["current_question_index"] >= len(state["questions"]):
        return jsonify({
            "status": "COMPLETED",
            "transcript": text,
            "action": "complete",
            "message": "Interview completed. Please call /generate-report to get the final summary."
        })

    next_question = state["questions"][state["current_question_index"]]

    return jsonify({
        "status": "COMPLETED",
        "transcript": text,
        "action": "continue",
        "next_question": next_question
    })


@app.route("/generate-report", methods=["POST"])
def generate_report():
    """
    Generate final report. If 'transcript' is provided in the body (old behavior),
    we still accept it. Otherwise we will use the stored SESSION_STORE transcript.
    """
    data = request.get_json() or {}
    external_transcript = data.get("transcript")

    # If client sends transcript string (old behavior), we can still use it
    if external_transcript:
        transcript_for_llm = external_transcript
    else:
        # Build transcript text from stored Q&A
        session_data = SESSION_STORE.get("active")
        if not session_data:
            return jsonify({"error": "No interview session data found"}), 400

        lines = []
        for qa in session_data["transcript"]:
            q = qa.get("question", "")
            a = qa.get("answer", "")
            if q:
                lines.append(f"Q: {q}")
            if a:
                lines.append(f"A: {a}")
        transcript_for_llm = "\n".join(lines)

    try:
        chain = get_report_generation_chain(model_id=MODEL_ID, region_name=REGION)
        report = chain.invoke({
            "transcript": transcript_for_llm
        })

        structured = SESSION_STORE.get("active", {}).get("structured", {})

        return jsonify({
            "summary": report.summary,
            "key_strengths": report.key_strengths,
            "areas_for_improvement": report.areas_for_improvement,
            "recommendation": report.recommendation,
            "structured": structured
        })

    except Exception as e:
        logger.error(f"Error generating report: {e}")
        return jsonify({"error": str(e)}), 500


# ============================
#           MAIN
# ============================

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
