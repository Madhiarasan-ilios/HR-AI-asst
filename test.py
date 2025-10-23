import os
import re
import boto3
import requests
import streamlit as st
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from docx.shared import Pt
from langchain_aws import ChatBedrock

# ================== ENV & AWS CONFIG ==================
load_dotenv()

AWS_ACCESS_KEY = os.getenv("access_key")
AWS_SECRET_KEY = os.getenv("secret_access_key")
AWS_REGION = os.getenv("region_name", "ap-south-1")
SERP_API_KEY = os.getenv("SERP_API_KEY")

session = boto3.Session(
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)

llm = ChatBedrock(
    model_id="anthropic.claude-3-sonnet-20240229-v1:0",
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region=AWS_REGION
)

# ================== HELPERS ==================

def clean_bedrock_response(response) -> str:
    """Extract only the content text from Claude response and clean formatting."""
    text = response if isinstance(response, str) else str(response)
    match = re.search(r'content="(.*?)"\s*(?:additional_kwargs|response_metadata|$)', text, re.DOTALL)
    clean = match.group(1) if match else text
    clean = clean.encode('utf-8').decode('unicode_escape')
    clean = re.sub(r'additional_kwargs=.*', '', clean)
    clean = re.sub(r'response_metadata=.*', '', clean)
    clean = re.sub(r'id=.*', '', clean)
    clean = re.sub(r'usage_metadata=.*', '', clean)
    return clean.strip().strip('"').strip("'").strip()


def ask_bedrock(prompt: str) -> str:
    """Invoke Bedrock model and clean the output."""
    try:
        response = llm.invoke(input=prompt)
        return clean_bedrock_response(response)
    except Exception as e:
        st.error(f"Bedrock error: {e}")
        return ""


def create_word_doc(content: str, filename: str) -> BytesIO:
    """Create a well-formatted Word document from JD text."""
    doc = Document()
    doc.add_heading("Job Description", level=1)

    sections = re.split(
        r'^(Job Description: AI Developer|Intro Section:|Objectives of this role:|Your tasks:|Required skills and qualifications:|Preferred skills and qualifications:)',
        content, flags=re.MULTILINE | re.IGNORECASE
    )

    for i in range(1, len(sections), 2):
        title = sections[i].strip()
        body = sections[i + 1].strip()

        if not body:
            continue

        if title.lower().startswith(('intro section', 'objectives', 'your tasks', 'required', 'preferred')):
            doc.add_heading(title, level=2)
        elif title.lower().startswith('job description'):
            intro_lines = body.split("\n\n")[0].split("\n")
            for line in intro_lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(8)
            continue

        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.lstrip().startswith(('•', '-', '*')):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(('•', '-', '*')):
                        cleaned_line = re.sub(r'^[•\-\*\s\t]+', '', line).strip()
                        if cleaned_line:
                            doc.add_paragraph(cleaned_line, style="List Bullet")
            else:
                p = doc.add_paragraph(block)
                p.paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gather_requirements():
    """UI for requirement gathering"""
    st.subheader("📋 Requirement Gathering")
    role = st.text_input("Role Title")
    experience = st.text_input("Experience (e.g., 3-5 years, Entry level)")
    domain = st.text_input("Domain/Industry (e.g., FinTech, HealthTech)")
    location = st.text_input("Job Location (e.g., Bengaluru, Remote, New York)")  # 🆕 Added field
    job_responsibilities = st.text_area("Job Responsibilities (projects or initiatives the candidate will work on)")
    company_values = st.text_area("Company Values & Culture")
    diversity = st.text_area("Diversity Considerations")

    return {
        "role": role,
        "experience": experience,
        "domain": domain,
        "location": location,  # 🆕 include in return dict
        "job_responsibilities": job_responsibilities,
        "company_values": company_values,
        "diversity": diversity
    }


def fetch_market_jds(role, num_results=5):
    """Fetch job postings using SERP API"""
    url = "https://serpapi.com/search"
    params = {
        "engine": "google_jobs",
        "q": role,
        "hl": "en",
        "num": num_results,
        "api_key": SERP_API_KEY
    }
    try:
        resp = requests.get(url, params=params)
        resp.raise_for_status()
        jobs = resp.json().get("jobs_results", [])
        return [job.get("description", "") for job in jobs]
    except Exception as e:
        st.error(f"Error fetching market JDs: {e}")
        return []


def extract_keywords_with_ai(postings):
    """Extract trending skills, phrases, bias terms"""
    text = "\n\n".join(postings[:5])
    prompt = f"""
    You are an HR analyst. Analyze the following job descriptions:

    {text}

    1. List trending technical and soft skills.
    2. Highlight frequently used phrases.
    3. Identify bias-prone or exclusive terms (e.g., 'rockstar', 'young').
    """
    return ask_bedrock(prompt)


def generate_job_description(requirements, analysis):
    """Generate structured JD using Claude, now includes Location."""
    prompt = f"""
    Draft an inclusive, engaging, SEO-optimized Job Description in the following template format.

    ============================
    Job Description: Template
    We are looking for a passionate {requirements['role']} to join our team at [Company X], located in {requirements['location']}.

    Intro Section:
    Provide 1–2 paragraphs introducing the role, responsibilities, and why this job is exciting.
    Explicitly mention the key project(s) or initiatives the new hire will be working on, based on:
    "{requirements['job_responsibilities']}"

    Objectives of this role:
    - List 4–5 key objectives directly connected to these projects and the role’s impact.

    Your tasks:
    - List 4–6 detailed tasks reflecting actual day-to-day work related to the provided project(s).
    - Ensure they align with the overall goals mentioned above.

    Required skills and qualifications:
    - 5–7 bullet points with mandatory skills, education, or certifications.
    - Include both technical and soft skills.
    - Include years of experience if applicable ({requirements['experience']}).

    Preferred skills and qualifications:
    - 3–5 bullet points with additional desirable or domain-specific skills.

    ============================

    Requirements from user:
    Role: {requirements['role']}
    Experience: {requirements['experience']}
    Domain: {requirements['domain']}
    Location: {requirements['location']}
    Company Values: {requirements['company_values']}
    Diversity Considerations: {requirements['diversity']}
    Key Project/Responsibilities: {requirements['job_responsibilities']}

    Market Insights:
    {analysis}

    Instructions:
    - Follow the template exactly.
    - Use inclusive, bias-free, gender-neutral language.
    - Naturally embed relevant keywords from market insights.
    - Include the location naturally in the introduction and conclusion.
    - End with a call-to-action encouraging candidates to apply.
    """
    return ask_bedrock(prompt)


def check_bias(jd_text):
    """Bias & inclusivity check"""
    prompt = f"""
    Review the following Job Description for biased, non-inclusive, or exclusionary language:

    {jd_text}

    Suggest improvements in bullet points.
    """
    return ask_bedrock(prompt)


# ================== STREAMLIT UI ==================
st.set_page_config(page_title="JD Agent - AWS", layout="wide")
st.title("🤖 Job Description AI Agent (AWS Bedrock)")

requirements = gather_requirements()

if st.button("Run JD Agent Pipeline"):
    if not requirements['role']:
        st.warning("Please fill at least the Role field.")
    else:
        postings = fetch_market_jds(requirements['role'])

        if postings:
            analysis = extract_keywords_with_ai(postings)
            st.subheader("📊 Market & Keyword Analysis")
            st.write(analysis)

            jd_text = generate_job_description(requirements, analysis)
            st.subheader("📝 Generated Job Description")
            st.text_area("Draft JD", jd_text, height=300)

            bias_feedback = check_bias(jd_text)
            st.subheader("✅ Bias & Inclusivity Feedback")
            st.write(bias_feedback)

            word_buffer = create_word_doc(jd_text, f"{requirements['role'].replace(' ', '_')}_JD.docx")
            st.download_button(
                label="⬇️ Download JD as Word (.docx)",
                data=word_buffer,
                file_name=f"{requirements['role'].replace(' ', '_')}_JD.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.info("No market postings found. JD will be generated without external comparison.")
